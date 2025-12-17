"""
Azure OpenAI service clients.

This module provides:
- AzureOpenAIClient: Base client for Azure OpenAI
- EmbeddingGenerator: Generate text embeddings
- TextGenerator: Generate text completions
- ResumeGenerator: AI-powered resume generation
"""
from io import BytesIO
import openai
from openai import AzureOpenAI
from config import Config
from core.rate_limiting import TokenUsageTracker, RateLimiter
    
# Check for optional dependencies
try:
    import httpx
    HTTPX_AVAILABLE = True
except ImportError:
    HTTPX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, black
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class AzureOpenAIClient:
    """Base Azure OpenAI client."""
    def __init__(self):
        # Clean endpoint to prevent double /openai path issues
        endpoint = Config.AZURE_OPENAI_ENDPOINT
        if endpoint:
            endpoint = endpoint.rstrip('/')
            if endpoint.endswith('/openai'):
                endpoint = endpoint[:-7]
        
        client_args = {
            "api_key": Config.AZURE_OPENAI_API_KEY,
            "api_version": Config.AZURE_OPENAI_API_VERSION,
            "azure_endpoint": endpoint
        }
        
        if HTTPX_AVAILABLE:
            # Create a custom http client that ignores SSL errors (sometimes needed in cloud envs)
            http_client = httpx.Client(verify=False)
            client_args["http_client"] = http_client
            
        self.client = AzureOpenAI(**client_args)
        self.token_tracker = TokenUsageTracker()
        self.rate_limiter = RateLimiter(max_calls=60, time_window=60)


class EmbeddingGenerator(AzureOpenAIClient):
    """Generate embeddings using Azure OpenAI."""
    def generate(self, text: str, model: str = None):
        model = model or Config.AZURE_OPENAI_EMBEDDING_DEPLOYMENT
        
        if not self.rate_limiter.allow_request():
            raise Exception("Rate limit exceeded")
        
        response = self.client.embeddings.create(
            input=text,
            model=model
        )
        
        self.token_tracker.add_usage(
            model=model,
            prompt_tokens=response.usage.prompt_tokens,
            completion_tokens=0
        )
        
        return response.data[0].embedding


class TextGenerator(AzureOpenAIClient):
    """Generate text using Azure OpenAI.
    
    This class provides text generation capabilities including:
    - General text generation
    - Resume generation tailored to job postings
    
    Flow for resume generation:
        ui/resume_tailor_page.py
          ↓
        services/azure_openai.py
          → TextGenerator.generate_resume()
          ↓
        modules/resume_generator/formatters.py
          → generate_docx_from_json()
    """
    
    def generate(self, prompt: str, model: str = None, **kwargs):
        """Generate text from a prompt.
        
        Args:
            prompt: The prompt to generate text from
            model: Optional model name (defaults to configured deployment)
            **kwargs: Additional arguments to pass to the API
            
        Returns:
            Generated text content
        """
        model = model or Config.AZURE_OPENAI_DEPLOYMENT
        
        if not self.rate_limiter.allow_request():
            raise Exception("Rate limit exceeded")
        
        response = self.client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            **kwargs
        )
        
        self.token_tracker.add_usage(
            model=model,
            prompt_tokens=response.usage.prompt_tokens,
            completion_tokens=response.usage.completion_tokens
        )
        
        return response.choices[0].message.content

    def generate_resume(self, user_profile: dict, job_posting: dict, 
                       raw_resume_text: str = None, model: str = None) -> dict:
        """Generate a tailored resume based on user profile and job posting.
        
        Uses the Context Sandwich approach for optimal results:
        1. System instructions for resume writing expertise
        2. Job posting details for targeting
        3. User profile and experience
        4. Optional raw resume text for additional context
        
        Args:
            user_profile: Dictionary with name, email, skills, experience, etc.
            job_posting: Dictionary with title, company, description, skills, etc.
            raw_resume_text: Optional original resume text for reference
            model: Optional model name (defaults to configured deployment)
            
        Returns:
            Structured resume data dictionary with header, summary, skills, 
            experience, education, and certifications sections.
            Returns None on error.
        """
        import json
        import re
        
        model = model or Config.AZURE_OPENAI_DEPLOYMENT
        
        if not self.rate_limiter.allow_request():
            raise Exception("Rate limit exceeded")
        
        system_instructions = """You are an expert resume writer with expertise in ATS optimization and career coaching.
Your task is to create a tailored resume by analyzing the job description and adapting the user's profile.
Return ONLY valid JSON - no markdown, no additional text, no code blocks."""

        job_description = f"""JOB POSTING TO MATCH:
Title: {job_posting.get('title', 'N/A')}
Company: {job_posting.get('company', 'N/A')}
Description: {job_posting.get('description', 'N/A')[:3000]}
Required Skills: {', '.join(job_posting.get('skills', [])[:10]) if job_posting.get('skills') else 'N/A'}"""

        structured_profile = f"""STRUCTURED PROFILE:
Name: {user_profile.get('name', 'N/A')}
Email: {user_profile.get('email', 'N/A')}
Phone: {user_profile.get('phone', 'N/A')}
Location: {user_profile.get('location', 'N/A')}
LinkedIn: {user_profile.get('linkedin', 'N/A')}
Summary: {user_profile.get('summary', 'N/A')}
Experience: {user_profile.get('experience', 'N/A')[:2000]}
Education: {user_profile.get('education', 'N/A')}
Skills: {user_profile.get('skills', 'N/A')}
Certifications: {user_profile.get('certifications', 'N/A')}"""

        raw_resume_section = ""
        if raw_resume_text:
            raw_resume_section = f"\n\nORIGINAL RESUME TEXT (for reference):\n{raw_resume_text[:2000]}"

        prompt = f"""{system_instructions}

{job_description}

{structured_profile}{raw_resume_section}

INSTRUCTIONS:
1. Analyze the job posting and identify key skills, technologies, and qualifications needed
2. Tailor the profile to match by:
   - Rewriting the summary to emphasize relevant experience
   - Highlighting skills that match job requirements
   - Rewriting experience bullet points to emphasize relevant achievements
   - Using keywords from the job description for ATS optimization
3. Focus on achievements and measurable results
4. Maintain accuracy - only use information from the provided profile

Return your response as a JSON object with this structure:
{{
  "header": {{
    "name": "Full Name",
    "title": "Professional Title (tailored to job)",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or empty string"
  }},
  "summary": "2-3 sentence professional summary tailored to the job",
  "skills_highlighted": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Date Range",
      "bullets": ["Achievement bullet 1...", "Achievement bullet 2..."]
    }}
  ],
  "education": "Education details",
  "certifications": "Certifications and achievements"
}}

Return ONLY the JSON object."""

        try:
            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_instructions},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.7,
                response_format={"type": "json_object"}
            )
            
            self.token_tracker.add_usage(
                model=model,
                prompt_tokens=response.usage.prompt_tokens,
                completion_tokens=response.usage.completion_tokens
            )
            
            content = response.choices[0].message.content
            
            # Parse JSON response
            try:
                content = content.strip()
                # Remove markdown code blocks if present
                if content.startswith("```"):
                    lines = content.split('\n')
                    content = '\n'.join(lines[1:-1]) if lines[-1].startswith('```') else '\n'.join(lines[1:])
                
                resume_data = json.loads(content)
                return resume_data
                
            except json.JSONDecodeError as e:
                # Try to extract JSON from the content
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    resume_data = json.loads(json_match.group())
                    return resume_data
                else:
                    print(f"Could not parse JSON response: {e}")
                    return None
                    
        except openai.APIConnectionError as e:
            print(f"❌ Connection error: {e}. Please check your AZURE_OPENAI_ENDPOINT.")
            return None
            
        except openai.AuthenticationError as e:
            print(f"❌ Authentication error: {e}. Please check your AZURE_OPENAI_API_KEY.")
            return None
            
        except Exception as e:
            print(f"Error generating resume: {e}")
            return None


# ============================================================================
# RESUME FORMATTERS
# ============================================================================

def set_cell_shading(cell, color: str):
    """Set background color for a table cell.
    
    Args:
        cell: DOCX table cell object
        color: Hex color string (without #)
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc, color: str = "2B5797"):
    """Add a horizontal line to the document.
    
    Args:
        doc: DOCX document object
        color: Hex color string for the line
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx_from_json(resume_data: dict, filename: str = "resume.docx"):
    """Generate a modern professional .docx file from structured resume JSON.
    
    Args:
        resume_data: Dictionary containing resume sections
        filename: Output filename (not used, returns BytesIO)
        
    Returns:
        BytesIO object containing the DOCX file, or None on error
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        PRIMARY_COLOR = RGBColor(43, 87, 151)
        SECONDARY_COLOR = RGBColor(80, 80, 80)
        ACCENT_COLOR = RGBColor(0, 120, 212)
        
        header = resume_data.get('header', {})
        
        # Name Header
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'].upper())
            name_run.font.size = Pt(24)
            name_run.font.bold = True
            name_run.font.color.rgb = PRIMARY_COLOR
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.paragraph_format.space_after = Pt(4)
        
        # Professional Title
        if header.get('title'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(header['title'])
            title_run.font.size = Pt(13)
            title_run.font.color.rgb = SECONDARY_COLOR
            title_run.font.italic = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(8)
        
        # Contact Info
        contact_items = []
        if header.get('email'):
            contact_items.append(f"✉ {header['email']}")
        if header.get('phone'):
            contact_items.append(f"📞 {header['phone']}")
        if header.get('location'):
            contact_items.append(f"📍 {header['location']}")
        if header.get('linkedin'):
            contact_items.append(f"💼 {header['linkedin']}")
        
        if contact_items:
            contact_para = doc.add_paragraph()
            contact_text = '  •  '.join(contact_items)
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = SECONDARY_COLOR
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(12)
        
        add_horizontal_line(doc, "2B5797")
        
        # Professional Summary
        if resume_data.get('summary'):
            summary_header = doc.add_paragraph()
            header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(resume_data['summary'])
            summary_run.font.size = Pt(10)
            summary_run.font.color.rgb = SECONDARY_COLOR
        
        # Key Skills
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            skills_header = doc.add_paragraph()
            header_run = skills_header.add_run('KEY SKILLS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            skills_para = doc.add_paragraph()
            for i, skill in enumerate(skills):
                skill_run = skills_para.add_run(f" {skill} ")
                skill_run.font.size = Pt(9)
                skill_run.font.color.rgb = PRIMARY_COLOR
                if i < len(skills) - 1:
                    separator = skills_para.add_run("  |  ")
                    separator.font.size = Pt(9)
                    separator.font.color.rgb = RGBColor(180, 180, 180)
        
        # Professional Experience
        experience = resume_data.get('experience', [])
        if experience:
            exp_header = doc.add_paragraph()
            header_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            for exp in experience:
                job_header = doc.add_paragraph()
                
                if exp.get('title'):
                    title_run = job_header.add_run(exp['title'])
                    title_run.font.size = Pt(11)
                    title_run.font.bold = True
                    title_run.font.color.rgb = RGBColor(50, 50, 50)
                
                if exp.get('company'):
                    company_run = job_header.add_run(f"  |  {exp['company']}")
                    company_run.font.size = Pt(10)
                    company_run.font.color.rgb = ACCENT_COLOR
                
                if exp.get('dates'):
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(exp['dates'])
                    date_run.font.size = Pt(9)
                    date_run.font.italic = True
                    date_run.font.color.rgb = SECONDARY_COLOR
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run("▸  ")
                        bullet_run.font.size = Pt(9)
                        bullet_run.font.color.rgb = ACCENT_COLOR
                        
                        text_run = bullet_para.add_run(bullet.strip())
                        text_run.font.size = Pt(10)
                        text_run.font.color.rgb = SECONDARY_COLOR
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        # Education
        if resume_data.get('education'):
            edu_header = doc.add_paragraph()
            header_run = edu_header.add_run('EDUCATION')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(resume_data['education'])
            edu_run.font.size = Pt(10)
            edu_run.font.color.rgb = SECONDARY_COLOR
        
        # Certifications
        if resume_data.get('certifications'):
            cert_header = doc.add_paragraph()
            header_run = cert_header.add_run('CERTIFICATIONS & ACHIEVEMENTS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(resume_data['certifications'])
            cert_run.font.size = Pt(10)
            cert_run.font.color.rgb = SECONDARY_COLOR
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        import streamlit as st
        st.error(f"Error generating DOCX: {e}")
        return None


def generate_pdf_from_json(resume_data: dict, filename: str = "resume.pdf"):
    """Generate a professional PDF file from structured resume JSON.
    
    Args:
        resume_data: Dictionary containing resume sections
        filename: Output filename (not used, returns BytesIO)
        
    Returns:
        BytesIO object containing the PDF file, or None on error
    """
    if not REPORTLAB_AVAILABLE:
        import streamlit as st
        st.error("PDF generation requires reportlab. Install with: pip install reportlab")
        return None
    
    try:
        pdf_io = BytesIO()
        doc = SimpleDocTemplate(
            pdf_io, 
            pagesize=letter,
            rightMargin=0.5*inch, 
            leftMargin=0.5*inch,
            topMargin=0.4*inch, 
            bottomMargin=0.4*inch
        )
        
        elements = []
        
        PRIMARY_COLOR = HexColor('#2B5797')
        SECONDARY_COLOR = HexColor('#505050')
        ACCENT_COLOR = HexColor('#0078D4')
        
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=PRIMARY_COLOR,
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=SECONDARY_COLOR,
            spaceAfter=4,
            alignment=TA_JUSTIFY
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=PRIMARY_COLOR,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        header = resume_data.get('header', {})
        
        if header.get('name'):
            elements.append(Paragraph(header['name'].upper(), name_style))
        
        contact_items = []
        if header.get('email'):
            contact_items.append(header['email'])
        if header.get('phone'):
            contact_items.append(header['phone'])
        if header.get('location'):
            contact_items.append(header['location'])
        
        if contact_items:
            contact_style = ParagraphStyle('ContactStyle', parent=styles['Normal'], 
                                          fontSize=9, textColor=SECONDARY_COLOR, alignment=TA_CENTER)
            elements.append(Paragraph('  •  '.join(contact_items), contact_style))
        
        elements.append(Spacer(1, 0.1*inch))
        elements.append(HRFlowable(width="100%", thickness=2, color=PRIMARY_COLOR, spaceAfter=0.1*inch))
        
        if resume_data.get('summary'):
            elements.append(Paragraph('PROFESSIONAL SUMMARY', section_header_style))
            elements.append(Paragraph(resume_data['summary'], body_style))
        
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            elements.append(Paragraph('KEY SKILLS', section_header_style))
            skills_text = '  |  '.join(skills)
            skills_style = ParagraphStyle('SkillsStyle', parent=styles['Normal'], 
                                         fontSize=9, textColor=PRIMARY_COLOR, alignment=TA_CENTER)
            elements.append(Paragraph(skills_text, skills_style))
        
        experience = resume_data.get('experience', [])
        if experience:
            elements.append(Paragraph('PROFESSIONAL EXPERIENCE', section_header_style))
            
            for exp in experience:
                if exp.get('title'):
                    job_style = ParagraphStyle('JobStyle', parent=styles['Normal'], 
                                              fontSize=11, textColor=black, fontName='Helvetica-Bold')
                    elements.append(Paragraph(exp['title'], job_style))
                
                if exp.get('company') or exp.get('dates'):
                    company_style = ParagraphStyle('CompanyStyle', parent=styles['Normal'], 
                                                  fontSize=10, textColor=ACCENT_COLOR)
                    company_text = f"{exp.get('company', '')}  |  {exp.get('dates', '')}"
                    elements.append(Paragraph(company_text, company_style))
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_style = ParagraphStyle('BulletStyle', parent=styles['Normal'], 
                                                     fontSize=10, textColor=SECONDARY_COLOR, leftIndent=15)
                        elements.append(Paragraph(f"▸  {bullet.strip()}", bullet_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        if resume_data.get('education'):
            elements.append(Paragraph('EDUCATION', section_header_style))
            elements.append(Paragraph(resume_data['education'], body_style))
        
        if resume_data.get('certifications'):
            elements.append(Paragraph('CERTIFICATIONS & ACHIEVEMENTS', section_header_style))
            elements.append(Paragraph(resume_data['certifications'], body_style))
        
        doc.build(elements)
        pdf_io.seek(0)
        return pdf_io
        
    except Exception as e:
        import streamlit as st
        st.error(f"Error generating PDF: {e}")
        return None


def format_resume_as_text(resume_data: dict) -> str:
    """Format structured resume JSON as plain text.
    
    Args:
        resume_data: Dictionary containing resume sections
        
    Returns:
        Formatted plain text resume
    """
    text = []
    
    header = resume_data.get('header', {})
    
    if header.get('name'):
        name = header['name'].upper()
        text.append("=" * 60)
        text.append(name.center(60))
        text.append("=" * 60)
        text.append("")
    
    if header.get('title'):
        text.append(header['title'].center(60))
        text.append("")
    
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
        text.append("-" * 60)
        text.append("")
    
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 25)
        text.append(resume_data['summary'])
        text.append("")
    
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 25)
        for i in range(0, len(skills), 4):
            row_skills = skills[i:i+4]
            text.append("  •  ".join(row_skills))
        text.append("")
    
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 25)
        for exp in experience:
            job_line = ""
            if exp.get('title'):
                job_line = exp['title']
            if exp.get('company'):
                job_line += f" | {exp['company']}"
            if job_line:
                text.append(job_line)
            
            if exp.get('dates'):
                text.append(f"    {exp['dates']}")
            
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet and bullet.strip():
                    text.append(f"    ▸ {bullet.strip()}")
            text.append("")
    
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 25)
        text.append(resume_data['education'])
        text.append("")
    
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & ACHIEVEMENTS")
        text.append("-" * 25)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)
