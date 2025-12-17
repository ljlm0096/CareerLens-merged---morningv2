# core/resume_generator.py
"""
Resume generation and formatting functionality.

This module provides:
- DOCX resume generation from structured JSON
- PDF resume generation from structured JSON
- Plain text resume formatting

Note: Heavy dependencies (docx, reportlab) are lazy-loaded to improve startup time.
"""

import streamlit as st
from io import BytesIO
from typing import Dict, Any

# Lazy imports for heavy modules - only load when needed
_docx_loaded = False
_Document = None
_Inches = None
_Pt = None
_RGBColor = None
_WD_ALIGN_PARAGRAPH = None
_qn = None
_OxmlElement = None


def _load_docx_modules():
    """Lazy load python-docx modules"""
    global _docx_loaded, _Document, _Inches, _Pt, _RGBColor, _WD_ALIGN_PARAGRAPH, _qn, _OxmlElement
    if not _docx_loaded:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        _Document = Document
        _Inches = Inches
        _Pt = Pt
        _RGBColor = RGBColor
        _WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        _qn = qn
        _OxmlElement = OxmlElement
        _docx_loaded = True
    return _Document, _Inches, _Pt, _RGBColor, _WD_ALIGN_PARAGRAPH, _qn, _OxmlElement


def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    _, _, _, _, _, qn, OxmlElement = _load_docx_modules()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc, color="2B5797"):
    """Add a horizontal line to the document"""
    _, _, Pt, _, _, qn, OxmlElement = _load_docx_modules()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Create a horizontal line using a bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Line thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx_from_json(resume_data: Dict[str, Any], filename: str = "resume.docx") -> BytesIO:
    """Generate a modern professional .docx file from structured resume JSON"""
    try:
        # Lazy load docx modules
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, _, _ = _load_docx_modules()
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Define colors
        PRIMARY_COLOR = RGBColor(43, 87, 151)  # Professional blue
        SECONDARY_COLOR = RGBColor(80, 80, 80)  # Dark gray
        ACCENT_COLOR = RGBColor(0, 120, 212)  # Bright blue for accents
        
        header = resume_data.get('header', {})
        
        # ===== NAME HEADER =====
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'].upper())
            name_run.font.size = Pt(24)
            name_run.font.bold = True
            name_run.font.color.rgb = PRIMARY_COLOR
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.paragraph_format.space_after = Pt(4)
        
        # ===== PROFESSIONAL TITLE =====
        if header.get('title'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(header['title'])
            title_run.font.size = Pt(13)
            title_run.font.color.rgb = SECONDARY_COLOR
            title_run.font.italic = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(8)
        
        # ===== CONTACT INFO BAR =====
        contact_items = []
        if header.get('email'):
            contact_items.append(f"✉ {header['email']}")
        if header.get('phone'):
            contact_items.append(f"📞 {header['phone']}")
        if header.get('location'):
            contact_items.append(f"📍 {header['location']}")
        if header.get('linkedin'):
            linkedin = header['linkedin']
            if 'linkedin.com' in linkedin:
                linkedin = linkedin.split('linkedin.com/')[-1].rstrip('/')
            contact_items.append(f"💼 {linkedin}")
        if header.get('portfolio'):
            contact_items.append(f"🌐 {header['portfolio']}")
        
        if contact_items:
            contact_para = doc.add_paragraph()
            contact_text = '  •  '.join(contact_items)
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = SECONDARY_COLOR
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(12)
        
        # Add decorative line
        add_horizontal_line(doc, "2B5797")
        
        # ===== PROFESSIONAL SUMMARY =====
        if resume_data.get('summary'):
            # Section header
            summary_header = doc.add_paragraph()
            header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            summary_header.paragraph_format.space_before = Pt(8)
            summary_header.paragraph_format.space_after = Pt(4)
            
            # Summary content
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(resume_data['summary'])
            summary_run.font.size = Pt(10)
            summary_run.font.color.rgb = SECONDARY_COLOR
            summary_para.paragraph_format.space_after = Pt(8)
        
        # ===== KEY SKILLS =====
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            skills_header = doc.add_paragraph()
            header_run = skills_header.add_run('KEY SKILLS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            skills_header.paragraph_format.space_before = Pt(8)
            skills_header.paragraph_format.space_after = Pt(4)
            
            # Create skill pills in a wrapped format
            skills_para = doc.add_paragraph()
            for i, skill in enumerate(skills):
                skill_run = skills_para.add_run(f" {skill} ")
                skill_run.font.size = Pt(9)
                skill_run.font.color.rgb = PRIMARY_COLOR
                if i < len(skills) - 1:
                    separator = skills_para.add_run("  |  ")
                    separator.font.size = Pt(9)
                    separator.font.color.rgb = RGBColor(180, 180, 180)
            skills_para.paragraph_format.space_after = Pt(8)
        
        # ===== PROFESSIONAL EXPERIENCE =====
        experience = resume_data.get('experience', [])
        if experience:
            exp_header = doc.add_paragraph()
            header_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            exp_header.paragraph_format.space_before = Pt(8)
            exp_header.paragraph_format.space_after = Pt(6)
            
            for exp in experience:
                # Job title and company on same line
                job_header = doc.add_paragraph()
                
                # Job title (bold)
                if exp.get('title'):
                    title_run = job_header.add_run(exp['title'])
                    title_run.font.size = Pt(11)
                    title_run.font.bold = True
                    title_run.font.color.rgb = RGBColor(50, 50, 50)
                
                # Company name
                if exp.get('company'):
                    company_run = job_header.add_run(f"  |  {exp['company']}")
                    company_run.font.size = Pt(10)
                    company_run.font.color.rgb = ACCENT_COLOR
                
                job_header.paragraph_format.space_after = Pt(0)
                
                # Date range (right-aligned style, but in new paragraph)
                if exp.get('dates'):
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(exp['dates'])
                    date_run.font.size = Pt(9)
                    date_run.font.italic = True
                    date_run.font.color.rgb = SECONDARY_COLOR
                    date_para.paragraph_format.space_after = Pt(4)
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_para = doc.add_paragraph()
                        # Custom bullet character
                        bullet_run = bullet_para.add_run("▸  ")
                        bullet_run.font.size = Pt(9)
                        bullet_run.font.color.rgb = ACCENT_COLOR
                        
                        text_run = bullet_para.add_run(bullet.strip())
                        text_run.font.size = Pt(10)
                        text_run.font.color.rgb = SECONDARY_COLOR
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        bullet_para.paragraph_format.space_after = Pt(2)
                
                # Add small spacing between experiences
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6)
        
        # ===== EDUCATION =====
        if resume_data.get('education'):
            edu_header = doc.add_paragraph()
            header_run = edu_header.add_run('EDUCATION')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            edu_header.paragraph_format.space_before = Pt(8)
            edu_header.paragraph_format.space_after = Pt(4)
            
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(resume_data['education'])
            edu_run.font.size = Pt(10)
            edu_run.font.color.rgb = SECONDARY_COLOR
            edu_para.paragraph_format.space_after = Pt(8)
        
        # ===== CERTIFICATIONS =====
        if resume_data.get('certifications'):
            cert_header = doc.add_paragraph()
            header_run = cert_header.add_run('CERTIFICATIONS & ACHIEVEMENTS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            cert_header.paragraph_format.space_before = Pt(8)
            cert_header.paragraph_format.space_after = Pt(4)
            
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(resume_data['certifications'])
            cert_run.font.size = Pt(10)
            cert_run.font.color.rgb = SECONDARY_COLOR
        
        # Save document
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None


def generate_pdf_from_json(resume_data: Dict[str, Any], filename: str = "resume.pdf") -> BytesIO:
    """Generate a modern professional PDF file from structured resume JSON"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        pdf_io = BytesIO()
        doc = SimpleDocTemplate(
            pdf_io, 
            pagesize=letter,
            rightMargin=0.5*inch, 
            leftMargin=0.5*inch,
            topMargin=0.4*inch, 
            bottomMargin=0.4*inch
        )
        
        elements = []
        
        # Define colors
        PRIMARY_COLOR = HexColor('#2B5797')  # Professional blue
        SECONDARY_COLOR = HexColor('#505050')  # Dark gray
        ACCENT_COLOR = HexColor('#0078D4')  # Bright blue
        LIGHT_GRAY = HexColor('#E8E8E8')
        
        # Define styles
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=PRIMARY_COLOR,
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=26
        )
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=SECONDARY_COLOR,
            spaceAfter=8,
            alignment=TA_CENTER,
            fontName='Helvetica-Oblique',
            leading=14
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=SECONDARY_COLOR,
            spaceAfter=6,
            alignment=TA_CENTER,
            leading=12
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=PRIMARY_COLOR,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            borderPadding=0,
            leading=14
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=SECONDARY_COLOR,
            spaceAfter=4,
            leading=13,
            alignment=TA_JUSTIFY
        )
        
        job_title_style = ParagraphStyle(
            'JobTitleStyle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=black,
            spaceAfter=0,
            fontName='Helvetica-Bold',
            leading=14
        )
        
        company_style = ParagraphStyle(
            'CompanyStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=ACCENT_COLOR,
            spaceAfter=2,
            leading=12
        )
        
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=SECONDARY_COLOR,
            spaceAfter=4,
            fontName='Helvetica-Oblique',
            leading=11
        )
        
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=SECONDARY_COLOR,
            spaceAfter=3,
            leftIndent=15,
            leading=13
        )
        
        skills_style = ParagraphStyle(
            'SkillsStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=PRIMARY_COLOR,
            spaceAfter=6,
            alignment=TA_CENTER,
            leading=14
        )
        
        # ===== HEADER SECTION =====
        header = resume_data.get('header', {})
        
        if header.get('name'):
            elements.append(Paragraph(header['name'].upper(), name_style))
        
        if header.get('title'):
            elements.append(Paragraph(header['title'], title_style))
        
        # Contact info
        contact_items = []
        if header.get('email'):
            contact_items.append(header['email'])
        if header.get('phone'):
            contact_items.append(header['phone'])
        if header.get('location'):
            contact_items.append(header['location'])
        if header.get('linkedin'):
            linkedin = header['linkedin']
            if 'linkedin.com/in/' in linkedin:
                linkedin = 'linkedin.com/in/' + linkedin.split('linkedin.com/in/')[-1].rstrip('/')
            elif 'linkedin.com' in linkedin:
                linkedin = linkedin.split('linkedin.com/')[-1].rstrip('/')
            contact_items.append(linkedin)
        if header.get('portfolio'):
            contact_items.append(header['portfolio'])
        
        if contact_items:
            contact_text = '  •  '.join(contact_items)
            elements.append(Paragraph(contact_text, contact_style))
        
        # Decorative line
        elements.append(Spacer(1, 0.1*inch))
        elements.append(HRFlowable(width="100%", thickness=2, color=PRIMARY_COLOR, spaceAfter=0.1*inch))
        
        # ===== PROFESSIONAL SUMMARY =====
        if resume_data.get('summary'):
            elements.append(Paragraph('PROFESSIONAL SUMMARY', section_header_style))
            elements.append(Paragraph(resume_data['summary'], body_style))
            elements.append(Spacer(1, 0.05*inch))
        
        # ===== KEY SKILLS =====
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            elements.append(Paragraph('KEY SKILLS', section_header_style))
            
            # Format skills with separators
            skills_text = '  |  '.join([f'<font color="#2B5797">{skill}</font>' for skill in skills])
            elements.append(Paragraph(skills_text, skills_style))
            elements.append(Spacer(1, 0.05*inch))
        
        # ===== PROFESSIONAL EXPERIENCE =====
        experience = resume_data.get('experience', [])
        if experience:
            elements.append(Paragraph('PROFESSIONAL EXPERIENCE', section_header_style))
            
            for exp in experience:
                # Job title
                if exp.get('title'):
                    elements.append(Paragraph(exp['title'], job_title_style))
                
                # Company and dates
                company_date_parts = []
                if exp.get('company'):
                    company_date_parts.append(f'<font color="#0078D4">{exp["company"]}</font>')
                if exp.get('dates'):
                    company_date_parts.append(f'<i>{exp["dates"]}</i>')
                
                if company_date_parts:
                    elements.append(Paragraph('  |  '.join(company_date_parts), company_style))
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_text = f'<font color="#0078D4">▸</font>  {bullet.strip()}'
                        elements.append(Paragraph(bullet_text, bullet_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        # ===== EDUCATION =====
        if resume_data.get('education'):
            elements.append(Paragraph('EDUCATION', section_header_style))
            elements.append(Paragraph(resume_data['education'], body_style))
            elements.append(Spacer(1, 0.05*inch))
        
        # ===== CERTIFICATIONS =====
        if resume_data.get('certifications'):
            elements.append(Paragraph('CERTIFICATIONS & ACHIEVEMENTS', section_header_style))
            elements.append(Paragraph(resume_data['certifications'], body_style))
        
        # Build PDF
        doc.build(elements)
        pdf_io.seek(0)
        return pdf_io
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None


def format_resume_as_text(resume_data: Dict[str, Any]) -> str:
    """Format structured resume JSON as plain text with modern formatting"""
    text = []
    
    header = resume_data.get('header', {})
    
    # Name header
    if header.get('name'):
        name = header['name'].upper()
        text.append("=" * 60)
        text.append(name.center(60))
        text.append("=" * 60)
        text.append("")
    
    # Professional title
    if header.get('title'):
        text.append(header['title'].center(60))
        text.append("")
    
    # Contact info
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    if header.get('linkedin'):
        contact.append(header['linkedin'])
    if header.get('portfolio'):
        contact.append(header['portfolio'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
        text.append("-" * 60)
        text.append("")
    
    # Professional Summary
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 25)
        text.append(resume_data['summary'])
        text.append("")
    
    # Key Skills
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 25)
        # Format skills in rows of 4
        for i in range(0, len(skills), 4):
            row_skills = skills[i:i+4]
            text.append("  •  ".join(row_skills))
        text.append("")
    
    # Professional Experience
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 25)
        for exp in experience:
            # Job header
            job_line = ""
            if exp.get('title'):
                job_line = exp['title']
            if exp.get('company'):
                job_line += f" | {exp['company']}"
            if job_line:
                text.append(job_line)
            
            if exp.get('dates'):
                text.append(f"    {exp['dates']}")
            
            # Bullets
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet and bullet.strip():
                    text.append(f"    ▸ {bullet.strip()}")
            text.append("")
    
    # Education
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 25)
        text.append(resume_data['education'])
        text.append("")
    
    # Certifications
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & ACHIEVEMENTS")
        text.append("-" * 25)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)


class ResumeGenerator:
    """Class-based interface for resume generation (optional wrapper)"""
    
    @staticmethod
    def generate_docx(resume_data: Dict[str, Any], filename: str = "resume.docx") -> BytesIO:
        """Generate DOCX resume from JSON data."""
        return generate_docx_from_json(resume_data, filename)
    
    @staticmethod
    def generate_pdf(resume_data: Dict[str, Any], filename: str = "resume.pdf") -> BytesIO:
        """Generate PDF resume from JSON data."""
        return generate_pdf_from_json(resume_data, filename)
    
    @staticmethod
    def format_as_text(resume_data: Dict[str, Any]) -> str:
        """Format resume as plain text."""
        return format_resume_as_text(resume_data)
